from __future__ import annotations

import argparse
import os
import re
import shutil
from pathlib import Path
from typing import Iterable

from docx import Document

try:
    from ragflow_sdk import RAGFlow
except ImportError:  # pragma: no cover
    RAGFlow = None


ROOT = Path(__file__).resolve().parents[2]
DATA_DIR = ROOT / "data"
PACKAGE_DIR = DATA_DIR / "示范景区公开资料包"
ALIAS_DIR = ROOT / ".playwright-mcp" / "kb_inspect"
OUTPUT_DIR = DATA_DIR / "ragflow_import"

STRUCTURED_ORIGINAL = PACKAGE_DIR / "灵山胜境 景点结构化数据集.docx"
GUIDE_ORIGINAL = PACKAGE_DIR / "灵山胜境：历史、文化、景点特色与个性化游览指南.docx"
STRUCTURED_ALIAS = ALIAS_DIR / "structured_spots.docx"
GUIDE_ALIAS = ALIAS_DIR / "guide.docx"

SECTION_TITLES = {
    "景区概况与千年历史渊源": "history",
    "小灵山的佛教缘起": "history",
    "祥符禅寺的千年兴衰": "history",
    "现代灵山胜境的崛起": "history",
    "核心文化内涵：佛教传承与艺术融合的典范": "culture",
    "佛教文化的深度传承": "culture",
    "传统艺术与现代科技的完美融合": "culture",
    "祈福文化的特色体验": "culture",
    "世界佛教文化的交流平台": "culture",
    "核心景点特色详解：佛教艺术的殿堂": "spots",
    "个性化游览路线推荐：深度体验灵山胜境": "route",
    "历史文化爱好者路线（6小时深度游）": "route",
    "自然风光爱好者路线（5小时全景游）": "route",
    "亲子家庭路线（4小时轻松游）": "route",
    "实用游览贴士：全方位保障你的灵山之旅": "tips",
    "门票与优惠政策": "tips",
    "最佳游览时间": "tips",
    "餐饮与住宿推荐": "tips",
    "其他实用建议": "tips",
}


def main() -> None:
    parser = argparse.ArgumentParser(description="Prepare and optionally upload RAGFlow knowledge base documents.")
    parser.add_argument("--upload", action="store_true", help="Upload prepared files to RAGFlow.")
    parser.add_argument("--replace", action="store_true", help="Delete prepared local output before regeneration.")
    args = parser.parse_args()

    if args.replace and OUTPUT_DIR.exists():
        shutil.rmtree(OUTPUT_DIR)

    structured_path = resolve_path(STRUCTURED_ORIGINAL, STRUCTURED_ALIAS)
    guide_path = resolve_path(GUIDE_ORIGINAL, GUIDE_ALIAS)

    spot_dir = OUTPUT_DIR / "scenic-spot-structured"
    history_dir = OUTPUT_DIR / "history-culture-route"
    spot_dir.mkdir(parents=True, exist_ok=True)
    history_dir.mkdir(parents=True, exist_ok=True)

    spot_files = build_structured_dataset(structured_path, spot_dir)
    history_files = build_guide_dataset(guide_path, history_dir)

    print(f"Prepared {len(spot_files)} scenic spot files in {spot_dir}")
    print(f"Prepared {len(history_files)} history/route files in {history_dir}")

    if args.upload:
        upload_to_ragflow(spot_dir, history_dir)


def resolve_path(original: Path, alias: Path) -> Path:
    if original.exists():
        return original
    if alias.exists():
        return alias
    raise FileNotFoundError(f"Source document not found: {original} or {alias}")


def build_structured_dataset(source: Path, output_dir: Path) -> list[Path]:
    doc = Document(str(source))
    table = doc.tables[0]
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    files: list[Path] = []

    for row in table.rows[1:]:
        item = dict(zip(headers, [cell.text.strip() for cell in row.cells]))
        if item.get("景区名称") != "灵山胜境":
            continue
        spot_id = item.get("景点ID", "unknown")
        spot_name = item.get("景点名称", "unknown")
        file_path = output_dir / safe_name(f"{spot_id}_{spot_name}.txt")
        file_path.write_text(render_structured_item(item), encoding="utf-8")
        files.append(file_path)
    return files


def render_structured_item(item: dict[str, str]) -> str:
    lines = [
        "kb_type: spot_structured",
        f"spot_id: {item.get('景点ID', '')}",
        f"spot_name: {item.get('景点名称', '')}",
        f"location_scope: {item.get('具体位置', '')}",
        f"has_open_info: {'yes' if item.get('演艺/开放信息') else 'no'}",
        "content_tag: spot_intro",
        "",
    ]
    for key in ["景区名称", "景点ID", "景点名称", "具体位置", "建筑/景观参数", "核心功能", "文化内涵", "详细介绍", "游玩亮点", "演艺/开放信息", "备注"]:
        lines.append(f"{key}: {item.get(key, '')}")
    return "\n".join(lines).strip() + "\n"


def build_guide_dataset(source: Path, output_dir: Path) -> list[Path]:
    doc = Document(str(source))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    sections: list[tuple[str, list[str], str]] = []
    current_title = "overview"
    current_type = "history"
    current_lines: list[str] = []

    for paragraph in paragraphs[1:]:
        if paragraph in SECTION_TITLES:
            if current_lines:
                sections.append((current_title, current_lines, current_type))
            current_title = paragraph
            current_type = SECTION_TITLES[paragraph]
            current_lines = []
        else:
            current_lines.append(paragraph)
    if current_lines:
        sections.append((current_title, current_lines, current_type))

    files: list[Path] = []
    for index, (title, lines, content_type) in enumerate(sections, start=1):
        for chunk_index, chunk in enumerate(chunk_lines(lines, 700), start=1):
            file_path = output_dir / safe_name(f"{index:02d}_{title}_{chunk_index}.txt")
            file_path.write_text(
                "\n".join([
                    "kb_type: history_route",
                    f"section_name: {title}",
                    f"route_type: {infer_route_type(title, chunk)}",
                    f"content_tag: {content_type}",
                    "",
                    chunk,
                    "",
                ]),
                encoding="utf-8",
            )
            files.append(file_path)
    return files


def chunk_lines(lines: Iterable[str], max_chars: int) -> list[str]:
    chunks: list[str] = []
    buffer: list[str] = []
    size = 0
    for line in lines:
        line_size = len(line)
        if buffer and size + line_size > max_chars:
            chunks.append("\n".join(buffer))
            buffer = [line]
            size = line_size
        else:
            buffer.append(line)
            size += line_size
    if buffer:
        chunks.append("\n".join(buffer))
    return chunks


def infer_route_type(title: str, chunk: str) -> str:
    if "亲子" in title or "亲子" in chunk:
        return "family"
    if "自然风光" in title:
        return "nature"
    if "历史文化" in title:
        return "culture"
    return "general"


def safe_name(name: str) -> str:
    return re.sub(r'[\\\\/:*?"<>|]+', "_", name)


def upload_to_ragflow(spot_dir: Path, history_dir: Path) -> None:
    if RAGFlow is None:
        raise RuntimeError("ragflow_sdk is not installed. Run: pip install ragflow-sdk")

    api_key = os.getenv("RAGFLOW_API_KEY")
    base_url = os.getenv("RAGFLOW_BASE_URL", "http://localhost:9380")
    if not api_key:
        raise RuntimeError("RAGFLOW_API_KEY is required for upload.")

    client = RAGFlow(api_key=api_key, base_url=base_url)
    datasets = {
        os.getenv("RAGFLOW_DATASET_SPOT_STRUCTURED", "scenic-spot-structured"): spot_dir,
        os.getenv("RAGFLOW_DATASET_HISTORY_ROUTE", "history-culture-route"): history_dir,
    }

    for dataset_name, directory in datasets.items():
        dataset = get_or_create_dataset(client, dataset_name)
        documents = []
        for file_path in sorted(directory.glob("*.txt")):
            documents.append({"display_name": file_path.name, "blob": file_path.read_bytes()})
        if documents:
            dataset.upload_documents(documents)
        print(f"Uploaded {len(documents)} documents to dataset: {dataset_name}")

    print("Upload complete. Create two RAGFlow chat assistants and bind them to the prepared datasets.")


def get_or_create_dataset(client: RAGFlow, dataset_name: str):
    existing = client.list_datasets(name=dataset_name)
    if existing:
        return existing[0]
    return client.create_dataset(name=dataset_name, chunk_method="one")


if __name__ == "__main__":
    main()
